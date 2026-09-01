from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(report_text, filename="Business_Report.docx"):
    """
    Generate a Word (.docx) version of the final business report.
    Mirrors reports/pdf_generator.py's line-based Markdown-lite
    conversion so both formats render the same content consistently.
    """

    document = Document()

    title = document.add_heading(
        "Development of Enterprise Workflow Platform "
        "with Decision Automation system",
        level=1,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_heading("Business Analysis Report", level=2)

    document.add_paragraph()

    for raw_line in report_text.split("\n"):
        line = raw_line.strip()

        if not line:
            continue

        line = line.replace("**", "")

        if line.endswith(":"):
            document.add_heading(line, level=3)
        else:
            paragraph = document.add_paragraph(line)
            paragraph.style.font.size = Pt(11)

    document.save(filename)

    return filename
